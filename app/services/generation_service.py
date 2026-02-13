"""
Document Generation Service - Creates PDF, DOCX, and CSV files.

Current Implementation:
    - PDF: WeasyPrint (markdown -> HTML -> PDF)
    - DOCX: python-docx (markdown -> DOCX with basic formatting)
    - CSV: Python csv module (dict list -> CSV bytes)
"""

from __future__ import annotations

import csv
import io
import logging
import re
from typing import Any, Dict, List, Optional

import markdown

logger = logging.getLogger("document_service.generation")


# Default CSS for PDF generation
DEFAULT_PDF_CSS = """
@page {
    size: letter;
    margin: 1in;
}

body {
    font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto,
                 "Helvetica Neue", Arial, sans-serif;
    font-size: 11pt;
    line-height: 1.6;
    color: #333;
}

h1 {
    font-size: 24pt;
    font-weight: 600;
    color: #1a1a1a;
    margin-top: 0;
    margin-bottom: 16pt;
    border-bottom: 2px solid #4f46e5;
    padding-bottom: 8pt;
}

h2 {
    font-size: 18pt;
    font-weight: 600;
    color: #1a1a1a;
    margin-top: 24pt;
    margin-bottom: 12pt;
}

h3 {
    font-size: 14pt;
    font-weight: 600;
    color: #333;
    margin-top: 18pt;
    margin-bottom: 8pt;
}

p {
    margin-bottom: 12pt;
}

ul, ol {
    margin-bottom: 12pt;
    padding-left: 24pt;
}

li {
    margin-bottom: 4pt;
}

code {
    font-family: "SF Mono", Monaco, "Courier New", monospace;
    font-size: 10pt;
    background-color: #f5f5f5;
    padding: 2pt 4pt;
    border-radius: 3pt;
}

pre {
    font-family: "SF Mono", Monaco, "Courier New", monospace;
    font-size: 10pt;
    background-color: #f5f5f5;
    padding: 12pt;
    border-radius: 4pt;
    overflow-x: auto;
    margin-bottom: 12pt;
}

blockquote {
    border-left: 4px solid #4f46e5;
    padding-left: 16pt;
    margin-left: 0;
    color: #555;
    font-style: italic;
}

table {
    width: 100%;
    border-collapse: collapse;
    margin-bottom: 12pt;
}

th, td {
    border: 1px solid #ddd;
    padding: 8pt;
    text-align: left;
}

th {
    background-color: #f5f5f5;
    font-weight: 600;
}

tr:nth-child(even) {
    background-color: #fafafa;
}

a {
    color: #4f46e5;
    text-decoration: none;
}

hr {
    border: none;
    border-top: 1px solid #ddd;
    margin: 24pt 0;
}

.title-page {
    text-align: center;
    padding-top: 200pt;
}

.title-page h1 {
    border-bottom: none;
    font-size: 32pt;
}

.metadata {
    color: #666;
    font-size: 10pt;
    margin-bottom: 24pt;
}
"""


class DocumentGenerationService:
    """
    Service for generating documents (PDF, DOCX, CSV) from content.

    This service is designed to be compartmentalized and replaceable.
    """

    def __init__(self):
        """Initialize the document generation service."""
        self._markdown_converter = markdown.Markdown(
            extensions=[
                "tables",
                "fenced_code",
                "toc",
                "nl2br",
            ]
        )

    # =========================================================================
    # PDF GENERATION
    # =========================================================================

    async def generate_pdf(
        self,
        content: str,
        title: Optional[str] = None,
        css: Optional[str] = None,
        include_title_page: bool = False,
    ) -> bytes:
        """
        Generate a PDF from markdown content.

        Args:
            content: Markdown content to convert
            title: Optional document title
            css: Optional custom CSS (uses default if not provided)
            include_title_page: If True, adds a title page at the beginning

        Returns:
            PDF file as bytes

        Raises:
            RuntimeError: If PDF generation fails
        """
        try:
            from weasyprint import CSS, HTML

            self._markdown_converter.reset()
            html_content = self._markdown_converter.convert(content)

            css_content = css or DEFAULT_PDF_CSS

            title_page_html = ""
            if include_title_page and title:
                title_page_html = f"""
                <div class="title-page">
                    <h1>{self._escape_html(title)}</h1>
                </div>
                <div style="page-break-after: always;"></div>
                """

            full_html = f"""
            <!DOCTYPE html>
            <html>
            <head>
                <meta charset="utf-8">
                <title>{self._escape_html(title or 'Document')}</title>
            </head>
            <body>
                {title_page_html}
                {html_content}
            </body>
            </html>
            """

            html_doc = HTML(string=full_html)
            css_doc = CSS(string=css_content)
            pdf_bytes = html_doc.write_pdf(stylesheets=[css_doc])

            logger.debug("Generated PDF: %d bytes", len(pdf_bytes))
            return pdf_bytes

        except ImportError as e:
            logger.error("WeasyPrint not installed: %s", e)
            raise RuntimeError(
                "PDF generation requires WeasyPrint. "
                "Install with: pip install weasyprint"
            ) from e
        except Exception as e:
            logger.error("PDF generation failed: %s", e)
            raise RuntimeError(f"PDF generation failed: {e}") from e

    # =========================================================================
    # DOCX GENERATION
    # =========================================================================

    async def generate_docx(
        self,
        content: str,
        title: Optional[str] = None,
    ) -> bytes:
        """
        Generate a DOCX file from markdown content.

        Args:
            content: Markdown content to convert
            title: Optional document title

        Returns:
            DOCX file as bytes

        Raises:
            RuntimeError: If DOCX generation fails
        """
        try:
            from docx import Document

            doc = Document()

            if title:
                doc.core_properties.title = title

            lines = content.split("\n")

            for line in lines:
                stripped = line.strip()

                if not stripped:
                    continue

                # Headers
                if stripped.startswith("### "):
                    doc.add_heading(stripped[4:], level=3)
                elif stripped.startswith("## "):
                    doc.add_heading(stripped[3:], level=2)
                elif stripped.startswith("# "):
                    doc.add_heading(stripped[2:], level=1)
                # Unordered list
                elif stripped.startswith("- ") or stripped.startswith("* "):
                    text = stripped[2:]
                    p = doc.add_paragraph(style="List Bullet")
                    self._add_formatted_text(p, text)
                # Ordered list
                elif re.match(r"^\d+\.\s", stripped):
                    text = re.sub(r"^\d+\.\s", "", stripped)
                    p = doc.add_paragraph(style="List Number")
                    self._add_formatted_text(p, text)
                # Horizontal rule
                elif stripped in ("---", "***", "___"):
                    doc.add_paragraph()
                # Regular paragraph
                else:
                    p = doc.add_paragraph()
                    self._add_formatted_text(p, stripped)

            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            docx_bytes = buffer.read()

            logger.debug("Generated DOCX: %d bytes", len(docx_bytes))
            return docx_bytes

        except ImportError as e:
            logger.error("python-docx not installed: %s", e)
            raise RuntimeError(
                "DOCX generation requires python-docx. "
                "Install with: pip install python-docx"
            ) from e
        except Exception as e:
            logger.error("DOCX generation failed: %s", e)
            raise RuntimeError(f"DOCX generation failed: {e}") from e

    def _add_formatted_text(self, paragraph, text: str):
        """Add text to a paragraph with basic markdown formatting (bold, italic)."""
        pattern = r"(\*\*\*(.+?)\*\*\*|\*\*(.+?)\*\*|\*(.+?)\*|([^*]+))"

        for match in re.finditer(pattern, text):
            if match.group(2):  # ***bold italic***
                run = paragraph.add_run(match.group(2))
                run.bold = True
                run.italic = True
            elif match.group(3):  # **bold**
                run = paragraph.add_run(match.group(3))
                run.bold = True
            elif match.group(4):  # *italic*
                run = paragraph.add_run(match.group(4))
                run.italic = True
            elif match.group(5):  # plain text
                paragraph.add_run(match.group(5))

    # =========================================================================
    # CSV GENERATION
    # =========================================================================

    async def generate_csv(
        self,
        data: List[Dict[str, Any]],
        columns: Optional[List[str]] = None,
        include_bom: bool = True,
    ) -> bytes:
        """
        Generate a CSV file from a list of dictionaries.

        Args:
            data: List of dictionaries (each dict is a row)
            columns: Optional list of column names (auto-detected if not provided)
            include_bom: If True, includes UTF-8 BOM for Excel compatibility

        Returns:
            CSV file as bytes (UTF-8 encoded)

        Raises:
            RuntimeError: If CSV generation fails
            ValueError: If data is empty and no columns provided
        """
        try:
            if not data and not columns:
                raise ValueError("Cannot generate CSV: no data or columns provided")

            if columns is None:
                if data:
                    all_keys = set()
                    for row in data:
                        all_keys.update(row.keys())
                    columns = sorted(all_keys)
                else:
                    columns = []

            buffer = io.StringIO()
            writer = csv.DictWriter(
                buffer,
                fieldnames=columns,
                extrasaction="ignore",
            )

            writer.writeheader()
            for row in data:
                writer.writerow(row)

            csv_content = buffer.getvalue()

            if include_bom:
                csv_bytes = b"\xef\xbb\xbf" + csv_content.encode("utf-8")
            else:
                csv_bytes = csv_content.encode("utf-8")

            logger.debug("Generated CSV: %d bytes, %d rows", len(csv_bytes), len(data))
            return csv_bytes

        except Exception as e:
            logger.error("CSV generation failed: %s", e)
            raise RuntimeError(f"CSV generation failed: {e}") from e

    # =========================================================================
    # UTILITY METHODS
    # =========================================================================

    def _escape_html(self, text: str) -> str:
        """Escape HTML special characters."""
        return (
            text.replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
            .replace('"', "&quot;")
            .replace("'", "&#39;")
        )

    def get_supported_formats(self) -> List[str]:
        """Get list of supported output formats."""
        return ["pdf", "docx", "csv"]


# Global service instance
document_generation_service = DocumentGenerationService()
